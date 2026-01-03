"""
Vector Store Manager
向量存储管理器实现
"""
import os
import time
import threading
from typing import Optional, Dict, Any, List
import logging
import tempfile

from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import PyPDFLoader, CSVLoader, TextLoader
from langchain_core.documents import Document as LangChainDocument

from interfaces.vector_store import VectorStoreInterface, EmbeddingInterface
from managers.cache_manager import CacheManager
from config import CHUNK_SIZE, CHUNK_OVERLAP

logger = logging.getLogger(__name__)


class WordDocumentLoader:
    """Word文档加载器（使用python-docx）"""
    
    def __init__(self, file_path: str):
        """
        初始化Word文档加载器
        
        Args:
            file_path: Word文档路径
        """
        self.file_path = file_path
    
    def load(self) -> List[LangChainDocument]:
        """加载Word文档"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(self.file_path)
            text_parts = []
            
            # 提取所有段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            # 合并所有文本
            full_text = "\n\n".join(text_parts)
            
            if not full_text.strip():
                logger.warning(f"Word document {self.file_path} appears to be empty")
                return []
            
            # 创建LangChain Document对象
            return [LangChainDocument(
                page_content=full_text,
                metadata={"source": self.file_path}
            )]
            
        except ImportError as import_err:
            error_msg = "python-docx library not installed. Please install it with: pip install python-docx"
            logger.error(error_msg)
            logger.error(f"Import error details: {import_err}")
            raise ImportError(error_msg) from import_err
        except Exception as e:
            logger.error(f"Error loading Word document {self.file_path}: {e}")
            import traceback
            logger.error(f"Word document loading traceback: {traceback.format_exc()}")
            raise


class ChromaVectorStoreManager(VectorStoreInterface):
    """ChromaDB向量存储管理器"""
    
    def __init__(self, embedding_interface: EmbeddingInterface):
        """
        初始化向量存储管理器
        
        Args:
            embedding_interface: 嵌入模型接口
        """
        self.embedding_interface = embedding_interface
        
        # 使用内存存储，不持久化
        self._vector_store = None
        self._vectorized_documents = []
        self._total_chunks = 0
        self._last_build_time = None
        
        self._lock = threading.RLock()
        logger.info("ChromaVectorStoreManager initialized (memory-only mode)")
    
    def get_store(self) -> Optional[Chroma]:
        """获取向量存储实例"""
        with self._lock:
            return self._vector_store
    
    def clear_store(self) -> bool:
        """清空向量存储和所有元数据
        
        Returns:
            清空是否成功
        """
        try:
            with self._lock:
                logger.info("Clearing vector store and metadata...")
                
                # 清空向量存储
                self._vector_store = None
                
                # 清空元数据
                self._vectorized_documents = []
                self._total_chunks = 0
                self._last_build_time = None
                
                logger.info("Vector store and metadata cleared successfully")
                return True
                
        except Exception as e:
            logger.error(f"Failed to clear vector store: {e}")
            return False
    
    def rebuild_store(self, documents_dir: str) -> bool:
        """
        重建向量存储 - 从文件系统（已废弃，保留用于接口兼容性）
        
        Args:
            documents_dir: 文档目录路径
        
        Returns:
            重建是否成功
        """
        logger.warning("rebuild_store is deprecated. Use rebuild_store_from_memory instead.")
        
        documents = self._load_documents(documents_dir)
        if not documents:
            logger.warning("No documents found to process")
            return False
                
        return self._build_vector_store_from_documents(documents)
    
    def rebuild_store_from_memory(self, in_memory_documents: Dict[str, bytes]) -> bool:
        """
        从内存文档重建向量存储
        
        Args:
            in_memory_documents: 内存中的文档字典 {filename: file_content_bytes}
        
        Returns:
            重建是否成功
        """
        if not in_memory_documents:
            logger.warning("No documents in memory to process")
            return False
        
        logger.info(f"Loading {len(in_memory_documents)} documents from memory")
        documents = self._load_documents_from_memory(in_memory_documents)
        if not documents:
            logger.warning("No documents loaded from memory")
            return False
        
        return self._build_vector_store_from_documents(documents)
    
    def _build_vector_store_from_documents(self, documents: List[Any]) -> bool:
        """
        从文档列表构建向量存储（公共方法，消除代码重复）
        
        Args:
            documents: 已加载的文档列表
        
        Returns:
            构建是否成功
        """
        with self._lock:
            try:
                logger.info("Starting vector store build...")
                
                # 1. 清除旧的内存存储
                self._vector_store = None
                self._vectorized_documents = []
                self._total_chunks = 0
                
                logger.info(f"Processing {len(documents)} documents")
                
                # 2. 处理文档为 chunks
                chunks = self._process_documents(documents)
                if not chunks:
                    logger.warning("No chunks generated from documents")
                    return False
                
                logger.info(f"Generated {len(chunks)} chunks")
                
                # 3. 获取嵌入模型
                embedding_model = self.embedding_interface.get_embeddings()
                if not embedding_model:
                    logger.error("Embedding model not available")
                    return False
                
                # 4. 创建向量存储（支持持久化）
                # 从环境变量读取持久化配置
                persist_dir = os.getenv("CHROMA_PERSIST_DIR", "")
                
                if persist_dir:
                    # 持久化模式
                    logger.info(f"Using persistent mode with directory: {persist_dir}")
                    os.makedirs(persist_dir, exist_ok=True)
                    self._vector_store = Chroma.from_documents(
                        documents=chunks,
                        embedding=embedding_model,
                        collection_name="documents",
                        persist_directory=persist_dir
                    )
                else:
                    # 纯内存模式（使用 EphemeralClient）
                    logger.info("Using ephemeral (in-memory) mode")
                    import chromadb
                    ephemeral_client = chromadb.EphemeralClient()
                    self._vector_store = Chroma(
                        client=ephemeral_client,
                        embedding_function=embedding_model,
                        collection_name="documents"
                    )
                    # 手动添加文档
                    self._vector_store.add_documents(documents=chunks)
                
                # 5. 记录向量化的文档信息
                document_set = set()
                for chunk in chunks:
                    if hasattr(chunk, 'metadata') and 'source' in chunk.metadata:
                        source = chunk.metadata['source']
                        filename = os.path.basename(source)
                        if filename:
                            document_set.add(filename)
                
                self._vectorized_documents = sorted(list(document_set))
                self._total_chunks = len(chunks)
                self._last_build_time = time.time()
                
                logger.info(f"Vector store built: {self._total_chunks} chunks, {len(self._vectorized_documents)} documents")
                
                return True
                
            except Exception as e:
                logger.error(f"Failed to build vector store: {e}")
                import traceback
                traceback.print_exc()
                return False
    
    def _load_documents_from_memory(self, in_memory_documents: Dict[str, bytes]) -> List[Any]:
        """从内存文档加载文档"""
        documents = []
        
        # 创建临时目录用于处理文件
        with tempfile.TemporaryDirectory() as temp_dir:
            for filename, file_content in in_memory_documents.items():
                try:
                    logger.info(f"Processing file from memory: {filename}")
                    
                    # 创建临时文件
                    temp_file_path = os.path.join(temp_dir, filename)
                    with open(temp_file_path, 'wb') as f:
                        f.write(file_content)
                    
                    # 使用LangChain加载器加载文档
                    if filename.lower().endswith('.pdf'):
                        loader = PyPDFLoader(temp_file_path)
                    elif filename.lower().endswith('.csv'):
                        loader = CSVLoader(temp_file_path)
                    elif filename.lower().endswith(('.txt', '.md')):
                        loader = TextLoader(temp_file_path)
                    elif filename.lower().endswith(('.docx', '.doc')):
                        loader = WordDocumentLoader(temp_file_path)
                    else:
                        logger.warning(f"Unsupported file type: {filename}")
                        continue
                    
                    docs = loader.load()
                    documents.extend(docs)
                    logger.info(f"Successfully loaded {len(docs)} pages from {filename}")
                    
                except Exception as e:
                    logger.error(f"Error loading {filename} from memory: {e}")
                    import traceback
                    logger.error(f"Traceback: {traceback.format_exc()}")
                    continue
        
        return documents
    
    def get_vectorized_documents(self) -> Dict[str, Any]:
        """获取已向量化的文档列表 - 内存模式"""
        with self._lock:
            if not self._vector_store:
                logger.info("No vector store available (memory mode)")
                return {"status": "success", "documents": [], "total_chunks": 0}
            
            logger.info(f"Returning cached vectorized documents: {len(self._vectorized_documents)} documents, {self._total_chunks} chunks")
            
            return {
                "status": "success", 
                "documents": self._vectorized_documents.copy(),
                "total_chunks": self._total_chunks
            }
    
    def is_available(self) -> bool:
        """检查向量存储是否可用"""
        with self._lock:
            return self._vector_store is not None
    
    def _load_documents(self, documents_dir: str) -> List[Any]:
        """加载文档"""
        documents = []
        
        if not os.path.exists(documents_dir):
            logger.error(f"Documents directory does not exist: {documents_dir}")
            return documents
        
        for filename in os.listdir(documents_dir):
            # 过滤系统文件
            if filename.startswith('.') or filename.startswith('~') or filename in ['.DS_Store', 'Thumbs.db', 'desktop.ini']:
                continue
            
            file_path = os.path.join(documents_dir, filename)
            if not os.path.isfile(file_path):
                continue
            
            try:
                logger.info(f"Processing file: {filename}")
                
                if filename.lower().endswith('.pdf'):
                    loader = PyPDFLoader(file_path)
                elif filename.lower().endswith('.csv'):
                    loader = CSVLoader(file_path)
                elif filename.lower().endswith(('.txt', '.md')):
                    loader = TextLoader(file_path)
                elif filename.lower().endswith(('.docx', '.doc')):
                    loader = WordDocumentLoader(file_path)
                else:
                    logger.warning(f"Unsupported file type: {filename}")
                    continue
                
                docs = loader.load()
                documents.extend(docs)
                logger.info(f"Successfully loaded {len(docs)} pages from {filename}")
                
            except Exception as e:
                logger.error(f"Error loading {filename}: {e}")
                continue
        
        return documents
    
    def _process_documents(self, documents: List[Any]) -> List[Any]:
        """处理文档为chunks"""
        try:
            from langchain_text_splitters import RecursiveCharacterTextSplitter
            
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=CHUNK_SIZE,
                chunk_overlap=CHUNK_OVERLAP,
                length_function=len,
            )
            
            chunks = text_splitter.split_documents(documents)
            logger.info(f"Generated {len(chunks)} chunks")
            
            # 打印前几个chunks的信息用于调试
            for i, chunk in enumerate(chunks[:5]):
                logger.debug(f"Chunk #{i+1} (length {len(chunk.page_content)}):")
                logger.debug(chunk.page_content[:200] + "..." if len(chunk.page_content) > 200 else chunk.page_content)
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing documents: {e}")
            return []
    

