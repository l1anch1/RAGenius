"""
Document Service
文档服务实现
"""
import os
import io
from typing import Dict, Any
import logging
import threading

from interfaces.services import DocumentServiceInterface
from interfaces.vector_store import VectorStoreInterface

logger = logging.getLogger(__name__)


class DocumentService(DocumentServiceInterface):
    """文档服务实现 - 内存模式"""
    
    def __init__(self, vector_store_manager: VectorStoreInterface):
        """
        初始化文档服务
        
        Args:
            vector_store_manager: 向量存储管理器
        """
        self.vector_store_manager = vector_store_manager
        
        # 内存中的文档存储：{filename: file_content_bytes}
        self._in_memory_documents: Dict[str, bytes] = {}
        self._lock = threading.RLock()
        
        logger.info("DocumentService initialized (memory mode)")
    
    def get_documents(self) -> Dict[str, Any]:
        """获取文档列表（从内存）"""
        try:
            with self._lock:
                documents = sorted(list(self._in_memory_documents.keys()))
                logger.info(f"Found {len(documents)} documents in memory: {documents}")
                return {"status": "success", "documents": documents}
            
        except Exception as e:
            logger.error(f"Failed to get documents: {e}")
            return {"status": "error", "message": str(e)}
    
    def get_vectorized_documents(self) -> Dict[str, Any]:
        """获取已向量化的文档列表"""
        try:
            return self.vector_store_manager.get_vectorized_documents()
        except Exception as e:
            logger.error(f"Failed to get vectorized documents: {e}")
            return {"status": "error", "message": str(e)}
    
    def upload_document(self, file) -> Dict[str, Any]:
        """上传文档到内存
        
        Args:
            file: Flask上传的文件对象
        
        Returns:
            上传结果
        """
        try:
            # 检查文件是否为空
            if not file or not file.filename:
                return {
                    "status": "error",
                    "message": "No file provided"
                }
            
            # 检查文件类型
            allowed_extensions = {'.pdf', '.txt', '.md', '.csv', '.docx', '.doc'}
            filename = file.filename
            file_ext = os.path.splitext(filename)[1].lower()
            
            if file_ext not in allowed_extensions:
                return {
                    "status": "error",
                    "message": f"Unsupported file type. Allowed types: {', '.join(sorted(allowed_extensions))}"
                }
            
            # 读取文件内容到内存
            # 重置文件指针到开头（以防之前被读取过）
            file.seek(0)
            file_content = file.read()
            
            # 验证文件内容不为空
            if not file_content:
                return {
                    "status": "error",
                    "message": "File is empty or could not be read"
                }
            
            with self._lock:
                # 检查文件是否已存在
                if filename in self._in_memory_documents:
                    return {
                        "status": "error",
                        "message": f"File '{filename}' already exists"
                    }
                
                # 保存到内存
                self._in_memory_documents[filename] = file_content
                logger.info(f"File uploaded to memory: {filename} ({len(file_content)} bytes)")
                logger.info(f"Total documents in memory: {len(self._in_memory_documents)}")
            
            return {
                "status": "success",
                "message": f"File '{filename}' uploaded successfully",
                "filename": filename
            }
            
        except Exception as e:
            logger.error(f"Failed to upload document: {e}")
            return {
                "status": "error",
                "message": f"Failed to upload document: {str(e)}"
            }
    
    def get_in_memory_documents(self) -> Dict[str, bytes]:
        """获取内存中的文档（供向量存储管理器使用）"""
        with self._lock:
            return self._in_memory_documents.copy()
    
    def get_document_preview(self, filename: str, max_length: int = 1000) -> Dict[str, Any]:
        """获取文档预览内容
        
        Args:
            filename: 文件名
            max_length: 最大预览字符数，默认1000
        
        Returns:
            预览结果
        """
        try:
            with self._lock:
                if filename not in self._in_memory_documents:
                    return {
                        "status": "error",
                        "message": f"Document '{filename}' not found"
                    }
                
                file_content = self._in_memory_documents[filename]
                file_ext = os.path.splitext(filename)[1].lower()
                
                # 根据文件类型处理预览
                preview_text = ""
                
                if file_ext in ['.txt', '.md', '.csv']:
                    # 文本文件直接解码
                    try:
                        preview_text = file_content.decode('utf-8')
                    except UnicodeDecodeError:
                        try:
                            preview_text = file_content.decode('gbk')
                        except:
                            preview_text = file_content.decode('latin-1')
                
                elif file_ext == '.pdf':
                    # PDF文件需要特殊处理
                    try:
                        import PyPDF2
                        pdf_file = io.BytesIO(file_content)
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                        
                        # 读取前几页
                        pages_to_read = min(3, len(pdf_reader.pages))
                        for page_num in range(pages_to_read):
                            page = pdf_reader.pages[page_num]
                            preview_text += page.extract_text() + "\n\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract PDF text: {e}")
                        preview_text = f"[PDF 文件，共 {len(file_content)} 字节]\n无法提取文本预览"
                
                elif file_ext in ['.docx', '.doc']:
                    # Word文档需要特殊处理
                    try:
                        from docx import Document
                        doc_file = io.BytesIO(file_content)
                        doc = Document(doc_file)
                        
                        # 读取前几段
                        paragraphs_to_read = min(10, len(doc.paragraphs))
                        for i in range(paragraphs_to_read):
                            preview_text += doc.paragraphs[i].text + "\n\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract Word text: {e}")
                        preview_text = f"[Word 文件，共 {len(file_content)} 字节]\n无法提取文本预览"
                
                else:
                    preview_text = f"[不支持的文件格式: {file_ext}]"
                
                # 截断到指定长度
                if len(preview_text) > max_length:
                    preview_text = preview_text[:max_length] + "\n\n... (内容已截断)"
                
                return {
                    "status": "success",
                    "filename": filename,
                    "preview": preview_text,
                    "size": len(file_content),
                    "type": file_ext
                }
        
        except Exception as e:
            logger.error(f"Failed to get document preview: {e}")
            return {
                "status": "error",
                "message": f"Failed to get document preview: {str(e)}"
            }
    
    def delete_document(self, filename: str) -> Dict[str, Any]:
        """删除单个文档
        
        Args:
            filename: 要删除的文件名
        
        Returns:
            删除结果
        """
        try:
            with self._lock:
                if filename not in self._in_memory_documents:
                    return {
                        "status": "error",
                        "message": f"Document '{filename}' not found"
                    }
                
                # 从内存中删除
                del self._in_memory_documents[filename]
                logger.info(f"Document deleted from memory: {filename}")
                logger.info(f"Remaining documents: {len(self._in_memory_documents)}")
                
                return {
                    "status": "success",
                    "message": f"Document '{filename}' deleted successfully",
                    "remaining_count": len(self._in_memory_documents)
                }
                
        except Exception as e:
            logger.error(f"Failed to delete document: {e}")
            return {
                "status": "error",
                "message": f"Failed to delete document: {str(e)}"
            }
    
    def clear_all_documents(self) -> Dict[str, Any]:
        """清空所有文档（包括内存和向量数据库）
        
        Returns:
            清空结果
        """
        try:
            with self._lock:
                # 清空内存中的文档
                doc_count = len(self._in_memory_documents)
                self._in_memory_documents.clear()
                logger.info(f"Cleared {doc_count} documents from memory")
                
                # 清空向量数据库
                self.vector_store_manager.clear_store()
                logger.info("Vector store cleared")
                
                return {
                    "status": "success",
                    "message": f"Successfully cleared {doc_count} documents and knowledge base",
                    "cleared_count": doc_count
                }
                
        except Exception as e:
            logger.error(f"Failed to clear documents: {e}")
            return {
                "status": "error",
                "message": f"Failed to clear documents: {str(e)}"
            }
    
    def rebuild_knowledge_base(self) -> Dict[str, Any]:
        """重建知识库（从内存中的文档）"""
        try:
            logger.info("Starting knowledge base rebuild from memory...")
            
            # 从内存获取文档
            in_memory_docs = self.get_in_memory_documents()
            
            logger.info(f"Found {len(in_memory_docs)} documents in memory: {list(in_memory_docs.keys())}")
            
            if not in_memory_docs:
                logger.warning("No documents in memory to rebuild")
                return {
                    "status": "error",
                    "message": "No documents in memory to rebuild. Please upload documents first."
                }
            
            # 使用内存文档重建向量存储
            success = self.vector_store_manager.rebuild_store_from_memory(in_memory_docs)
            
            if success:
                logger.info("Knowledge base rebuild completed successfully")
                return {
                    "status": "success",
                    "message": "Knowledge base rebuilt successfully"
                }
            else:
                logger.error("Knowledge base rebuild failed")
                import traceback
                logger.error(f"Rebuild traceback: {traceback.format_exc()}")
                return {
                    "status": "error",
                    "message": "Failed to rebuild knowledge base. Please check the logs for details."
                }
                
        except Exception as e:
            logger.error(f"Failed to rebuild knowledge base: {e}")
            import traceback
            logger.error(f"Exception traceback: {traceback.format_exc()}")
            return {
                "status": "error",
                "message": f"Failed to rebuild knowledge base: {str(e)}"
            }
